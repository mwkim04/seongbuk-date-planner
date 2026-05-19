import math
import os
import random
from io import BytesIO
from pathlib import Path
from typing import Dict, List, Tuple
from urllib.parse import quote

import folium
import pandas as pd
import requests
import streamlit as st
from docx import Document
from dotenv import load_dotenv
from streamlit_folium import st_folium

APP_DIR = Path(__file__).resolve().parent
load_dotenv(APP_DIR / ".env")

APP_VERSION = "final_cloud_secrets_fast_api"
PINK = "#ff4f9a"
HOT_PINK = "#ff2f86"
LIGHT_PINK = "#fff2f7"
BORDER = "#ffd0e2"
BLACK = "#111111"
GRAY = "#6b7280"

st.set_page_config(
    page_title="성북구 데이트 코스 자동 추천",
    page_icon="💗",
    layout="wide",
    initial_sidebar_state="expanded",
)

st.markdown(
    f"""
<style>
:root {{ --primary-color: {PINK}; }}
* {{ font-family: "Pretendard", -apple-system, BlinkMacSystemFont, "Segoe UI", "Apple SD Gothic Neo", "Noto Sans KR", sans-serif !important; }}
.stApp {{ background:#ffffff !important; color:{BLACK} !important; }}
header {{ visibility:hidden; height:0; }}
[data-testid="collapsedControl"], [data-testid="stSidebarCollapseButton"], button[title="Close sidebar"], button[title="Open sidebar"] {{ display:none !important; }}
.block-container {{ max-width:1320px; padding-top:2rem; padding-bottom:3rem; }}
section[data-testid="stSidebar"] {{ background:#f4f6fa !important; border-right:1px solid #e5e7eb; }}
section[data-testid="stSidebar"] > div:first-child {{ padding-top:0.1rem !important; }}
[data-testid="InputInstructions"], div[data-testid="stTextInput"] [data-testid="InputInstructions"] {{ display:none !important; visibility:hidden !important; height:0 !important; }}
section[data-testid="stSidebar"] * {{ color:{BLACK} !important; }}
.stMultiSelect [data-baseweb="tag"], div[data-baseweb="tag"] {{ background:{PINK} !important; color:#fff !important; border-radius:9px !important; }}
.stMultiSelect [data-baseweb="tag"] span, div[data-baseweb="tag"] span {{ color:#fff !important; }}
.stMultiSelect [data-baseweb="tag"] svg, div[data-baseweb="tag"] svg {{ color:#fff !important; fill:#fff !important; }}
input[type="range"] {{ accent-color:{PINK} !important; }}
input[type="checkbox"] {{ accent-color:{PINK} !important; }}
.stSlider [data-testid="stTickBar"] {{ background:{PINK} !important; }}
div[role="slider"] {{ background:{PINK} !important; border-color:{PINK} !important; }}
div[data-testid="stSlider"] [data-baseweb="slider"] div[aria-valuenow] {{ background:{PINK} !important; }}
/* Streamlit/BaseWeb 기본 빨강 포인트를 핑크로 강제 */
[data-testid="stSidebar"] [data-baseweb="tag"] {{ background-color:{PINK} !important; border-color:{PINK} !important; }}
[data-testid="stSidebar"] [data-baseweb="tag"] * {{ color:#fff !important; fill:#fff !important; }}
[data-testid="stSidebar"] [data-baseweb="slider"] div {{ border-color:{PINK} !important; }}
[data-testid="stSidebar"] [data-baseweb="slider"] [style*="background"] {{ background-color:{PINK} !important; }}
[data-testid="stSidebar"] div[role="slider"] {{ background-color:{PINK} !important; border-color:{PINK} !important; }}
[data-testid="stSidebar"] svg {{ color:inherit; }}
.hero {{
  position:relative; overflow:hidden;
  background:#ffffff; color:{BLACK};
  border:1px solid #f3d7e4; border-radius:28px;
  padding:36px 42px; margin:6px 0 24px 0;
  box-shadow:0 18px 42px rgba(255,79,154,.10), 0 6px 18px rgba(17,17,17,.05);
}}
.hero::before {{
  content:""; position:absolute; left:0; top:0; bottom:0; width:10px;
  background:linear-gradient(180deg,{HOT_PINK},{PINK});
}}
.hero::after {{
  content:""; position:absolute; right:-80px; top:-95px; width:260px; height:260px;
  background:radial-gradient(circle, rgba(255,79,154,.20) 0%, rgba(255,79,154,.07) 42%, rgba(255,255,255,0) 70%);
  pointer-events:none;
}}
.hero-inner {{ position:relative; z-index:1; display:flex; align-items:center; gap:22px; }}
.hero-icon {{
  width:76px; height:76px; border-radius:24px;
  display:flex; align-items:center; justify-content:center;
  background:linear-gradient(135deg,#fff0f7,#ffe0ed);
  border:1px solid #ffd0e2;
  box-shadow:0 10px 24px rgba(255,79,154,.18);
  font-size:36px; flex:0 0 auto;
}}
.hero-copy {{ flex:1; }}
.hero-kicker {{
  display:inline-flex; align-items:center; gap:8px;
  background:{LIGHT_PINK}; color:{HOT_PINK};
  border:1px solid {BORDER}; border-radius:999px;
  padding:6px 12px; margin-bottom:12px;
  font-size:13px; font-weight:900;
}}
.hero h1 {{ color:{BLACK} !important; font-size:42px; line-height:1.15; margin:0 0 12px 0; font-weight:950; letter-spacing:-.04em; }}
.hero p {{ color:#3f3f46 !important; margin:0; font-size:16px; font-weight:700; }}
.hero-highlight {{ color:{HOT_PINK}; }}
.flow {{ background:{LIGHT_PINK}; border:1px solid {BORDER}; color:{BLACK}; padding:16px 18px; border-radius:16px; font-weight:850; margin:12px 0 18px; }}
.notice {{ background:#fff7fb; border-left:5px solid {PINK}; padding:14px 16px; border-radius:14px; margin:14px 0 20px; color:{BLACK}; }}
.metric-card {{ background:#fff; border:1px solid #eee; border-radius:18px; padding:16px 18px; box-shadow:0 8px 24px rgba(0,0,0,.045); }}
.metric-title {{ color:{GRAY}; font-size:13px; font-weight:850; }}
.metric-value {{ color:{BLACK}; font-size:28px; font-weight:950; margin-top:4px; }}
.route-card {{ background:#fff; border:1px solid {BORDER}; border-radius:22px; padding:22px 24px; margin:20px 0 16px; box-shadow:0 10px 28px rgba(255,79,154,.08); }}
.route-card h2 {{ color:{BLACK}; margin:0 0 12px 0; font-size:34px; font-weight:950; }}
.stat-row {{ display:flex; gap:12px; flex-wrap:wrap; margin:14px 0 8px; }}
.stat {{ background:#fff; border:1px solid {BORDER}; border-left:6px solid {PINK}; border-radius:16px; padding:14px 18px; min-width:160px; }}
.stat .label {{ color:{GRAY}; font-size:13px; font-weight:850; }}
.stat .num {{ color:{BLACK}; font-size:30px; font-weight:950; margin-top:2px; }}
.pill {{ display:inline-block; padding:6px 10px; border-radius:999px; background:{LIGHT_PINK}; color:{HOT_PINK}; font-weight:900; font-size:12px; margin-right:6px; }}
div.stButton > button[kind="primary"], div.stFormSubmitButton > button {{
  background:linear-gradient(90deg,{HOT_PINK},{PINK}) !important; color:white !important; border:0 !important;
  border-radius:16px !important; font-weight:950 !important; padding:.95rem 1.1rem !important; font-size:18px !important;
  box-shadow:0 10px 24px rgba(255,79,154,.32) !important;
}}
div.stButton > button:not([kind="primary"]) {{
  background:#fff !important; color:{BLACK} !important; border:1px solid #e5e7eb !important; border-radius:14px !important; font-weight:800 !important;
}}
div.stDownloadButton > button {{
  background:#fff !important; color:{HOT_PINK} !important; border:1px solid {BORDER} !important; border-radius:14px !important; font-weight:900 !important;
}}
.leaflet-div-icon {{ background:transparent !important; border:none !important; }}
.num-marker {{
  width:30px; height:30px; border-radius:50%; background:{PINK}; color:white; display:flex; align-items:center; justify-content:center;
  font-weight:950; font-size:15px; border:3px solid #fff; box-shadow:0 3px 10px rgba(0,0,0,.24);
}}

/* 사이드바 포인트 색상 핑크 고정 */
section[data-testid="stSidebar"] div[data-baseweb="tag"] {{
  background-color:#ff4f9a !important;
  border-color:#ff4f9a !important;
  color:#ffffff !important;
}}
section[data-testid="stSidebar"] div[data-baseweb="tag"] * {{
  color:#ffffff !important;
  fill:#ffffff !important;
}}
section[data-testid="stSidebar"] input[type="checkbox"] {{
  accent-color:#ff4f9a !important;
}}

/* 체크박스는 박스만 핑크로 표시하고, 글자/라벨 배경은 유지 */
section[data-testid="stSidebar"] [data-testid="stCheckbox"] label {{
  background: transparent !important;
}}
section[data-testid="stSidebar"] [data-testid="stCheckbox"] label > div:first-child {{
  background: transparent !important;
}}
section[data-testid="stSidebar"] [data-testid="stCheckbox"] input[type="checkbox"] {{
  accent-color:#ff4f9a !important;
}}
section[data-testid="stSidebar"] div[role="slider"] {{
  background:#ff4f9a !important;
  border-color:#ff4f9a !important;
}}
section[data-testid="stSidebar"] [data-baseweb="slider"] [class*="InnerThumb"] {{
  background:#ff4f9a !important;
}}
section[data-testid="stSidebar"] [data-baseweb="slider"] div {{
  border-color:#ff4f9a !important;
}}

</style>
""",
    unsafe_allow_html=True,
)

REGION_KEYWORDS = {
    "안암동": ["안암동", "고려대", "고려대학교"],
    "성신여대": ["성신여대", "성신여대입구", "동선동"],
    "한성대입구": ["한성대입구", "삼선동", "한성대학교"],
    "혜화": ["혜화", "혜화역", "대학로"],
    "길음": ["길음", "길음역"],
    "정릉": ["정릉", "정릉동", "정릉천"],
    "성북구 전체": ["성북구", "성북동"],
}

REGION_CENTERS = {
    "안암동": (37.5864, 127.0290, 1.7),
    "성신여대": (37.5928, 127.0165, 1.4),
    "한성대입구": (37.5885, 127.0061, 1.3),
    "혜화": (37.5820, 127.0019, 1.4),
    "길음": (37.6034, 127.0248, 1.6),
    "정릉": (37.6083, 127.0098, 1.8),
}

SEARCH_KEYWORDS = {
    "식당": ["맛집", "식당", "한식", "중식", "일식", "양식", "파스타", "분식", "고기집", "돈까스", "소바", "라멘", "초밥"],
    "카페": ["카페", "디저트", "베이커리", "커피"],
    "포토부스": ["인생네컷", "포토이즘", "하루필름", "포토부스", "셀프사진관"],
    "활동": ["보드게임카페", "노래방", "영화관", "볼링장", "오락실", "방탈출"],
    "산책": ["공원", "산책", "전시", "미술관", "문화거리", "낙산공원", "정릉천", "성북천"],
    "술집": ["술집", "이자카야", "맥주", "호프", "포차", "와인바", "칵테일"],
}

# Streamlit Cloud에서는 많은 API 호출이 무한로딩처럼 보일 수 있어 검색량을 제한함.
MAX_REGION_TERMS = 2
MAX_KEYWORDS_PER_GROUP = 6
MAX_KAKAO_PAGES = 1

SUSPICIOUS_WORDS = [
    "무료선택", "폐업", "임시", "테스트", "삭제", "광고", "부동산", "공인중개사", "분양", "매매",
    "화장실", "개방화장실", "공중화장실", "지하철역", "전철역", "버스정류장", "정류장",
    "주짓수", "태권도", "합기도", "체육관", "헬스장",
    # 데이트 장소로 부적합한 생활/소매/미용 계열
    "준오헤어", "헤어", "미용실", "살롱", "네일", "피부관리", "왁싱", "마사지", "에스테틱",
    "올리브영", "랄라블라", "롭스", "화장품", "드럭스토어", "편집샵", "의류", "옷가게", "신발", "안경", "문구점"
]
EXCLUDE_CATEGORIES = [
    "부동산", "학원", "병원", "약국", "은행", "주차장", "편의점", "마트", "관공서", "학교", "회사", "기업",
    "화장실", "공중화장실", "교통", "지하철", "전철", "버스정류장", "정류장", "체육", "무술", "주짓수", "태권도",
    "교회", "성당", "사찰",
    # 카카오 검색 결과에 섞이는 비데이트 업종 제거
    "미용", "헤어", "뷰티", "네일", "피부", "마사지", "화장품", "생활", "쇼핑", "소매", "패션", "의류", "잡화", "문구", "안경"
]


def get_api_key() -> str:
    """로컬(.env)과 Streamlit Cloud(Secrets) 모두에서 카카오 REST API 키를 읽음."""
    try:
        key = st.secrets.get("KAKAO_REST_API_KEY", "")
        if key:
            return str(key).strip()
    except Exception:
        pass
    return os.getenv("KAKAO_REST_API_KEY", "").strip()


@st.cache_data(ttl=3600, show_spinner=False)
def kakao_search_cached(query: str, api_key: str, page: int = 1) -> Dict:
    url = "https://dapi.kakao.com/v2/local/search/keyword.json"
    headers = {"Authorization": f"KakaoAK {api_key}"}
    params = {"query": query, "page": page, "size": 15}
    response = requests.get(url, headers=headers, params=params, timeout=12)
    if response.status_code == 401:
        raise RuntimeError("카카오 REST API 키 인증 실패: Streamlit Secrets 또는 .env의 REST API 키를 확인하세요.")
    if response.status_code == 429:
        raise RuntimeError("카카오 API 호출 한도 초과: 잠시 후 다시 시도하세요.")
    response.raise_for_status()
    return response.json()


def is_non_date_place(name: str, category_name: str, address: str = "") -> bool:
    """카카오 API 검색 결과 중 데이트 코스 장소로 보기 어려운 업종을 제거."""
    text = f"{name} {category_name} {address}"
    bad_words = [
        "준오헤어", "헤어", "미용실", "살롱", "네일", "피부", "왁싱", "마사지", "에스테틱",
        "올리브영", "랄라블라", "롭스", "화장품", "드럭스토어", "편집샵", "의류", "옷가게", "신발", "안경", "문구",
        "병원", "의원", "약국", "부동산", "공인중개사", "주차장", "은행", "편의점", "마트",
        "교회", "성당", "사찰", "절", "학교", "학원", "회사", "기업", "체육관", "주짓수", "태권도", "헬스장",
        "지하철역", "전철역", "버스정류장", "정류장", "화장실"
    ]
    return any(w in text for w in bad_words)


def classify_place(name: str, category_name: str, requested_group: str) -> str:
    text = f"{name} {category_name}"
    if any(x in text for x in ["이자카야", "포차", "호프", "맥주", "와인바", "칵테일", "술집", "막걸리", "전집", "할머니맥주"]):
        return "술집"
    if any(x in text for x in ["인생네컷", "포토이즘", "하루필름", "포토부스", "셀프사진", "사진관"]):
        return "포토부스"
    if any(x in text for x in ["카페", "커피", "디저트", "베이커리", "케이크", "마카롱", "빵"]):
        return "카페"
    if any(x in text for x in ["공원", "산책로", "둘레길", "전시", "미술관", "박물관", "문화거리", "정릉천", "성북천"]) or name.endswith("공원") or "낙산공원" in name:
        return "산책"
    if any(x in text for x in ["보드게임", "노래방", "영화관", "CGV", "볼링", "오락", "방탈출", "PC방"]):
        return "활동"
    if any(x in category_name for x in ["음식점"]):
        return "식당"
    # 검색어 때문에 끼어든 무관한 장소는 요청 분류로 억지 편입하지 않음
    return "기타"


def cost_category_and_price(name: str, category_name: str, place_type: str) -> Tuple[str, int]:
    text = f"{name} {category_name}"
    if place_type == "포토부스":
        return "포토부스", 4000
    if place_type == "산책":
        if any(x in text for x in ["전시", "미술관", "박물관"]):
            return "전시/미술관", 8000
        return "공원/산책 (무료)", 0
    if place_type == "활동":
        if "노래방" in text:
            return "노래방", 8000
        if "보드게임" in text:
            return "보드게임카페", 7000
        if any(x in text for x in ["영화", "CGV"]):
            return "영화관", 15000
        if "볼링" in text:
            return "볼링장", 12000
        if "방탈출" in text:
            return "방탈출", 22000
        return "오락/활동", 8000
    if place_type == "카페":
        if any(x in text for x in ["디저트", "케이크", "마카롱", "빙수"]):
            return "디저트 카페", 8000
        if any(x in text for x in ["베이커리", "빵"]):
            return "베이커리 카페", 8500
        return "일반 카페", 6000
    if place_type == "술집":
        if "이자카야" in text:
            return "이자카야", 25000
        if any(x in text for x in ["와인", "칵테일"]):
            return "와인바/칵테일바", 30000
        if any(x in text for x in ["포차", "호프", "맥주"]):
            return "포차/호프", 17000
        return "일반 술집", 18000
    if any(x in text for x in ["중식", "중국", "짜장", "짬뽕", "마라"]):
        return "중식", 12000
    if any(x in text for x in ["일식", "초밥", "스시", "라멘", "돈카츠", "돈까스", "소바", "우동"]):
        return "일식", 14000
    if any(x in text for x in ["파스타", "스테이크", "양식", "피자", "리조또"]):
        return "양식", 16000
    if any(x in text for x in ["고기", "삼겹살", "갈비", "구이"]):
        return "고기/구이", 22000
    if any(x in text for x in ["분식", "김밥", "떡볶이", "라면"]):
        return "분식", 8000
    if any(x in text for x in ["국밥", "찌개", "백반", "한식"]):
        return "한식", 11000
    return "일반 식당", 12000


def mood_tags_for(name: str, category_name: str, place_type: str) -> List[str]:
    text = f"{name} {category_name} {place_type}"
    tags = set()
    if place_type in ["카페", "산책", "포토부스"]:
        tags.update(["감성", "조용한"])
    if place_type in ["활동", "술집"]:
        tags.update(["활동적"])
    if any(x in text for x in ["한옥", "레트로", "시장", "다방"]):
        tags.update(["레트로", "감성"])
    if any(x in text for x in ["포차", "호프", "맥주", "노래방", "볼링", "오락", "보드게임"]):
        tags.update(["활동적"])
    if any(x in text for x in ["분식", "김밥", "국밥", "백반", "버거", "마라", "돈까스", "소바"]):
        tags.update(["가성비"])
    if any(x in text for x in ["와인", "칵테일", "파스타", "스테이크", "디저트", "베이커리"]):
        tags.update(["감성"])
    if not tags:
        tags.add("감성")
    return sorted(tags)


def parse_terms(text: str) -> List[str]:
    return [t.strip() for t in str(text).replace('，', ',').replace('/', ',').split(',') if t.strip()]


def row_matches_terms(row: Dict, terms: List[str]) -> bool:
    if not terms:
        return True
    haystack = ' '.join(str(row.get(k, '')) for k in ['name', 'type', 'price_category', 'category_name', 'address', 'mood_tags'])
    return any(term in haystack for term in terms)


def is_in_selected_region(region: str, address: str, lat: float, lon: float) -> bool:
    if "서울" not in address:
        return False
    if region == "혜화":
        if "종로구" not in address:
            return False
    elif region == "성북구 전체":
        if "성북구" not in address:
            return False
        return True
    else:
        if "성북구" not in address:
            return False
    info = REGION_CENTERS.get(region)
    if info:
        center_lat, center_lon, radius_km = info
        if haversine_km((center_lat, center_lon), (lat, lon)) > radius_km:
            return False
    return True

def collect_places(region: str, include_bar: bool, api_key: str, include_terms: List[str] | None = None) -> pd.DataFrame:
    groups = ["식당", "카페", "포토부스", "활동", "산책"] + (["술집"] if include_bar else [])
    include_terms = include_terms or []
    rows = []
    for requested_group in groups:
        for region_term in REGION_KEYWORDS.get(region, [region])[:MAX_REGION_TERMS]:
            for keyword in SEARCH_KEYWORDS[requested_group][:MAX_KEYWORDS_PER_GROUP]:
                query = f"{region_term} {keyword}"
                for page in range(1, MAX_KAKAO_PAGES + 1):
                    data = kakao_search_cached(query, api_key, page)
                    for item in data.get("documents", []):
                        name = (item.get("place_name") or "").strip()
                        address = item.get("road_address_name") or item.get("address_name") or ""
                        category_name = item.get("category_name") or ""
                        url = item.get("place_url") or ""
                        if not name or not address or not url:
                            continue
                        try:
                            lat = float(item.get("y"))
                            lon = float(item.get("x"))
                        except Exception:
                            continue
                        if not is_in_selected_region(region, address, lat, lon):
                            continue
                        merged = f"{name} {category_name} {address}"
                        if any(word in merged for word in SUSPICIOUS_WORDS + EXCLUDE_CATEGORIES):
                            continue
                        if is_non_date_place(name, category_name, address):
                            continue
                        place_type = classify_place(name, category_name, requested_group)
                        if place_type == "기타":
                            continue
                        if place_type == "술집" and not include_bar:
                            continue
                        price_cat, cost = cost_category_and_price(name, category_name, place_type)
                        tags = mood_tags_for(name, category_name, place_type)
                        rows.append({
                            "name": name,
                            "type": place_type,
                            "price_category": price_cat,
                            "category_name": category_name,
                            "address": address,
                            "phone": item.get("phone", ""),
                            "lat": lat,
                            "lon": lon,
                            "url": url,
                            "cost": int(cost),
                            "mood_tags": ",".join(tags),
                            "rating": "카카오맵 확인",
                            "open_status": "카카오맵 확인 필요",
                            "source": "Kakao API",
                        })
                    if data.get("meta", {}).get("is_end", True):
                        break
    # 사용자가 반드시 포함하고 싶은 키워드는 별도로 한 번 더 검색해 후보군에 넣음
    for term in include_terms:
        for region_term in REGION_KEYWORDS.get(region, [region])[:MAX_REGION_TERMS]:
            query = f"{region_term} {term}"
            for page in range(1, MAX_KAKAO_PAGES + 1):
                data = kakao_search_cached(query, api_key, page)
                for item in data.get("documents", []):
                    name = (item.get("place_name") or "").strip()
                    address = item.get("road_address_name") or item.get("address_name") or ""
                    category_name = item.get("category_name") or ""
                    url = item.get("place_url") or ""
                    if not name or not address or not url:
                        continue
                    try:
                        lat = float(item.get("y"))
                        lon = float(item.get("x"))
                    except Exception:
                        continue
                    if not is_in_selected_region(region, address, lat, lon):
                        continue
                    merged = f"{name} {category_name} {address}"
                    if any(word in merged for word in SUSPICIOUS_WORDS + EXCLUDE_CATEGORIES):
                        continue
                    if is_non_date_place(name, category_name, address):
                        continue
                    place_type = classify_place(name, category_name, "식당")
                    if place_type == "기타":
                        continue
                    if place_type == "술집" and not include_bar:
                        continue
                    price_cat, cost = cost_category_and_price(name, category_name, place_type)
                    tags = mood_tags_for(name, category_name, place_type)
                    rows.append({
                        "name": name,
                        "type": place_type,
                        "price_category": price_cat,
                        "category_name": category_name,
                        "address": address,
                        "phone": item.get("phone", ""),
                        "lat": lat,
                        "lon": lon,
                        "url": url,
                        "cost": int(cost),
                        "mood_tags": ",".join(tags),
                        "rating": "카카오맵 확인",
                        "open_status": "카카오맵 확인 필요",
                        "source": "Kakao API",
                    })
                if data.get("meta", {}).get("is_end", True):
                    break

    if not rows:
        return pd.DataFrame()
    return pd.DataFrame(rows).drop_duplicates(subset=["name", "address"]).reset_index(drop=True)


def add_scores(df: pd.DataFrame, moods: List[str], include_terms: List[str] | None = None) -> pd.DataFrame:
    df = df.copy()
    include_terms = include_terms or []

    def score_row(row):
        tags = set(str(row["mood_tags"]).split(","))
        mood_score = sum(1 for mood in moods if mood in tags)
        keyword_bonus = 8 if row_matches_terms(row.to_dict(), include_terms) else 0
        return mood_score + keyword_bonus + random.random() * 0.25

    df["score"] = df.apply(score_row, axis=1)
    return df.reset_index(drop=True)


def haversine_km(a: Tuple[float, float], b: Tuple[float, float]) -> float:
    lat1, lon1 = a
    lat2, lon2 = b
    radius = 6371.0
    phi1 = math.radians(lat1)
    phi2 = math.radians(lat2)
    dphi = math.radians(lat2 - lat1)
    dlambda = math.radians(lon2 - lon1)
    h = math.sin(dphi / 2) ** 2 + math.cos(phi1) * math.cos(phi2) * math.sin(dlambda / 2) ** 2
    return 2 * radius * math.atan2(math.sqrt(h), math.sqrt(1 - h))


def estimate_minutes(distance_km: float, mode: str) -> int:
    speed = {"도보": 4.2, "대중교통": 13.0, "대중교통+도보": 9.0, "자동차": 18.0}.get(mode, 4.2)
    base = {"도보": 0, "대중교통": 8, "대중교통+도보": 6, "자동차": 5}.get(mode, 0)
    return max(1, round(distance_km / speed * 60 + base))


def direction_url(a: Dict, b: Dict) -> str:
    return f"https://map.kakao.com/link/from/{quote(a['name'])},{a['lat']},{a['lon']}/to/{quote(b['name'])},{b['lat']},{b['lon']}"


def sequence_templates(place_count: int, include_bar: bool, budget: int) -> List[List[str]]:
    """데이트 흐름을 자연스럽게 유지하면서 조건별로 다양한 순서를 제공."""
    if place_count <= 1:
        return [["식당"], ["카페"], ["산책"]]

    if include_bar:
        # 술집은 보통 뒤쪽에 배치. 저예산에서도 결과가 나오도록 무료/저가 활동 조합을 함께 제공.
        base = {
            2: [["카페", "술집"], ["산책", "술집"], ["식당", "술집"]],
            3: [["카페", "산책", "술집"], ["식당", "산책", "술집"], ["식당", "카페", "술집"], ["산책", "카페", "술집"]],
            4: [
                ["카페", "산책", "포토부스", "술집"],
                ["산책", "카페", "포토부스", "술집"],
                ["식당", "카페", "산책", "술집"],
                ["식당", "포토부스", "카페", "술집"],
                ["카페", "식당", "산책", "술집"],
            ],
            5: [
                ["식당", "카페", "포토부스", "산책", "술집"],
                ["카페", "산책", "식당", "포토부스", "술집"],
                ["산책", "카페", "식당", "활동", "술집"],
                ["식당", "활동", "카페", "포토부스", "술집"],
            ],
        }
    else:
        base = {
            2: [["식당", "카페"], ["카페", "산책"], ["산책", "카페"], ["카페", "포토부스"]],
            3: [["식당", "카페", "여가"], ["카페", "산책", "식당"], ["산책", "카페", "식당"], ["식당", "포토부스", "카페"]],
            4: [["식당", "카페", "포토부스", "산책"], ["카페", "산책", "식당", "포토부스"], ["산책", "카페", "식당", "활동"], ["식당", "활동", "카페", "산책"], ["카페", "식당", "포토부스", "산책"]],
            5: [["식당", "카페", "포토부스", "활동", "산책"], ["카페", "산책", "식당", "포토부스", "활동"], ["산책", "카페", "식당", "활동", "포토부스"], ["식당", "활동", "카페", "포토부스", "산책"]],
        }

    templates = list(base.get(place_count, base[5]))
    random.shuffle(templates)
    return templates


def sequence_for(place_count: int, include_bar: bool, budget: int) -> List[str]:
    return random.choice(sequence_templates(place_count, include_bar, budget))


def weighted_pick_from_candidates(cand: pd.DataFrame) -> Dict | None:
    """상위 후보 중에서 무작위로 선택해 같은 조건에서도 코스가 반복되지 않도록 함."""
    if cand.empty:
        return None
    cand = cand.copy().sort_values("rank", ascending=False).head(min(14, len(cand)))
    weights = [max(1.0, len(cand) - idx) for idx in range(len(cand))]
    return cand.sample(1, weights=weights, random_state=random.randint(0, 2_000_000_000)).iloc[0].to_dict()


def _slot_candidates(df: pd.DataFrame, slot: str) -> pd.DataFrame:
    if slot == "여가":
        return df[df["type"].isin(["포토부스", "활동", "산책"])].copy()
    return df[df["type"] == slot].copy()


def choose_candidate_bulletproof(
    df: pd.DataFrame,
    slot: str,
    used_names: set,
    used_leisure_types: set,
    prev: Dict | None,
    strict: bool = True,
    include_terms: List[str] | None = None,
    prefer_terms: bool = False,
) -> Dict | None:
    include_terms = include_terms or []
    cand = _slot_candidates(df, slot)

    # 같은 여가 유형이 반복되는 것을 줄임. 단 후보가 완전히 없으면 같은 유형 재사용보다 해당 코스는 포기.
    if slot == "여가" and strict:
        narrowed = cand[~cand["type"].isin(used_leisure_types)]
        if not narrowed.empty:
            cand = narrowed

    # 이름 중복은 피함. 후보가 부족하더라도 전혀 다른 분류로 억지 대체하지 않음.
    cand = cand[~cand["name"].isin(used_names)]
    if cand.empty:
        return None

    # 반드시 포함 키워드가 아직 충족되지 않았고 현재 슬롯에서 매칭 후보가 있으면 그 후보를 우선 선택.
    if prefer_terms and include_terms:
        term_mask = cand.apply(lambda r: row_matches_terms(r.to_dict(), include_terms), axis=1)
        if term_mask.any():
            cand = cand[term_mask]

    if prev:
        cand["near"] = cand.apply(lambda r: haversine_km((prev["lat"], prev["lon"]), (r["lat"], r["lon"])), axis=1)
    else:
        cand["near"] = 0.0

    cand["rank"] = cand["score"] * 2.2 - cand["near"] * 0.75 + [random.random() * 2.8 for _ in range(len(cand))]
    return weighted_pick_from_candidates(cand)


def _route_from_sequence(
    df: pd.DataFrame,
    seq: List[str],
    budget_limit: float,
    mode: str,
    max_dist: float,
    allow_over: bool,
    include_bar: bool,
    strict: bool,
    include_terms: List[str] | None = None,
) -> Dict | None:
    """한 개의 코스를 빠르게 생성. 오래 도는 무한대기 방지용."""
    include_terms = include_terms or []
    used_names, used_leisure_types = set(), set()
    places: List[Dict] = []
    prev = None

    for slot in seq:
        matched_already = any(row_matches_terms(p, include_terms) for p in places) if include_terms else True
        p = choose_candidate_bulletproof(
            df=df,
            slot=slot,
            used_names=used_names,
            used_leisure_types=used_leisure_types,
            prev=prev,
            strict=strict,
            include_terms=include_terms,
            prefer_terms=not matched_already,
        )
        if not p:
            return None
        if p["type"] == "술집" and not include_bar:
            return None
        places.append(p)
        used_names.add(p["name"])
        if p["type"] in ["포토부스", "활동", "산책"]:
            used_leisure_types.add(p["type"])
        prev = p

    if include_terms and not any(row_matches_terms(p, include_terms) for p in places):
        return None

    segments = []
    total_distance = 0.0
    for a, b in zip(places, places[1:]):
        dist = haversine_km((a["lat"], a["lon"]), (b["lat"], b["lon"]))
        total_distance += dist
        segments.append({
            "from": a["name"],
            "to": b["name"],
            "mode": mode,
            "distance_km": dist,
            "minutes": estimate_minutes(dist, mode),
            "guide": "도보 이동 예상" if mode == "도보" else "카카오맵에서 상세 노선 확인",
            "url": direction_url(a, b),
        })

    total_cost = sum(int(p["cost"]) for p in places)
    final_budget_limit = budget_limit * (1.15 if allow_over else 1.0)
    if total_cost > final_budget_limit:
        return None
    if total_distance > max_dist:
        return None

    return {
        "places": places,
        "segments": segments,
        "distance_km": total_distance,
        "minutes": sum(s["minutes"] for s in segments),
        "cost": total_cost,
        "budget_ok": total_cost <= budget_limit,
        "mode": mode,
    }


def build_routes(df: pd.DataFrame, budget: int, mode: str, include_bar: bool, place_count: int, max_distance: float, allow_over_budget: bool, include_terms: List[str] | None = None) -> List[Dict]:
    """요청 조건으로 최대 5개 추천안을 빠르게 생성. 술집 체크 시 오래 멈추는 현상을 방지."""
    include_terms = include_terms or []
    if include_terms and not df.apply(lambda r: row_matches_terms(r.to_dict(), include_terms), axis=1).any():
        raise RuntimeError(f"반드시 포함할 키워드({', '.join(include_terms)})와 일치하는 장소를 찾지 못했습니다. 키워드를 바꾸거나 지역을 넓혀보세요.")

    # 필요한 큰 분류 후보가 아예 없으면 바로 안내하고 멈춤.
    templates = sequence_templates(place_count, include_bar, budget)
    required_types = set()
    for seq in templates:
        for slot in seq:
            if slot == "여가":
                required_types.update(["포토부스", "활동", "산책"])
            else:
                required_types.add(slot)
    if include_bar:
        required_types.add("술집")
    missing = []
    for t in required_types:
        if t in ["포토부스", "활동", "산책"]:
            # 여가 세부 타입은 모두 없어도 되므로 개별 missing으로 보지 않음
            continue
        if df[df["type"] == t].empty:
            missing.append(t)
    if missing:
        raise RuntimeError(f"{', '.join(missing)} 후보가 부족합니다. 지역을 넓히거나 조건을 완화해보세요.")

    routes: List[Dict] = []
    seen_keys = set()

    def add_route(r: Dict | None):
        if not r:
            return
        key = tuple(p["name"] for p in r["places"])
        if key in seen_keys:
            return
        seen_keys.add(key)
        routes.append(r)

    # 1차: 엄격 조건으로 빠르게 생성
    for strict in [True, False]:
        for seq in templates:
            for _ in range(80):
                if len(routes) >= 5:
                    return routes[:5]
                add_route(_route_from_sequence(df, seq, budget, mode, max_distance, allow_over_budget, include_bar, strict, include_terms))

    # 2차: 예산 초과 허용 체크가 켜졌으면 거리만 조금 완화
    if len(routes) < 5 and allow_over_budget:
        for seq in templates:
            for _ in range(60):
                if len(routes) >= 5:
                    return routes[:5]
                add_route(_route_from_sequence(df, seq, budget, mode, max_distance * 1.25, True, include_bar, False, include_terms))

    return routes[:5]


# 기존의 오래 도는 매칭 함수는 남겨두지 않고 빠른 생성 함수로 대체함.

def sort_routes(routes: List[Dict], sort_by: str) -> List[Dict]:
    if not routes:
        return []
    if sort_by == "총 이동시간 오름차순":
        return sorted(routes, key=lambda r: r["minutes"])
    if sort_by == "총 이동시간 내림차순":
        return sorted(routes, key=lambda r: r["minutes"], reverse=True)
    if sort_by == "1인당 예상비용 오름차순":
        return sorted(routes, key=lambda r: r["cost"])
    if sort_by == "1인당 예상비용 내림차순":
        return sorted(routes, key=lambda r: r["cost"], reverse=True)
    if sort_by == "전체거리 오름차순":
        return sorted(routes, key=lambda r: r["distance_km"])
    if sort_by == "전체거리 내림차순":
        return sorted(routes, key=lambda r: r["distance_km"], reverse=True)
    return routes


def make_map(route: Dict):
    places = route["places"]
    center = [sum(p["lat"] for p in places) / len(places), sum(p["lon"] for p in places) / len(places)]
    m = folium.Map(location=center, zoom_start=15, tiles="OpenStreetMap", control_scale=True)
    coords = []
    for idx, p in enumerate(places, start=1):
        coords.append([p["lat"], p["lon"]])
        html = f"""
        <div style="
            width:44px;height:44px;border-radius:50% 50% 50% 0;
            background:linear-gradient(135deg,#ff7ab8 0%,#ff2f86 100%);
            transform:rotate(-45deg);
            border:4px solid #fff;
            box-shadow:0 5px 14px rgba(0,0,0,.28);
            display:flex;align-items:center;justify-content:center;">
            <div style="
                transform:rotate(45deg);
                color:#fff;font-weight:950;font-size:22px;
                line-height:1;text-shadow:0 2px 4px rgba(0,0,0,.35);">
                {idx}
            </div>
        </div>
        """
        popup_html = f"""
        <div style="font-family:Pretendard, sans-serif; min-width:180px;">
            <b style="font-size:15px;">{idx}. {p['name']}</b><br>
            <span style="color:#ff2f86;font-weight:700;">{p['type']}</span><br>
            <span>{p['address']}</span><br>
            <a href="{p['url']}" target="_blank">카카오맵 열기</a>
        </div>
        """
        folium.Marker(
            location=[p["lat"], p["lon"]],
            tooltip=f"{idx}. {p['name']}",
            popup=folium.Popup(popup_html, max_width=260),
            icon=folium.DivIcon(icon_size=(44, 44), icon_anchor=(22, 42), html=html),
        ).add_to(m)
    if len(coords) > 1:
        folium.PolyLine(coords, color="#4b5563", weight=9, opacity=0.25).add_to(m)
        folium.PolyLine(coords, color="#ffffff", weight=6, opacity=0.9).add_to(m)
        folium.PolyLine(coords, color=PINK, weight=4, opacity=0.95).add_to(m)
    try:
        m.fit_bounds(coords, padding=(28, 28))
    except Exception:
        pass
    return m


def create_docx_bytes(routes: List[Dict], meta: Dict) -> bytes:
    """실사용 가능한 데이트 플래너 DOCX 생성.

    화면 결과표를 그대로 옮기는 대신, 데이트 당일 휴대폰/인쇄물로 보기 좋은
    요약표, 타임라인, 장소 카드, 이동 정보, 예산 요약, 체크리스트를 포함한다.
    """
    from docx.enum.section import WD_SECTION
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.4)
    section.left_margin = Cm(1.45)
    section.right_margin = Cm(1.45)

    PRIMARY = "FF4F9A"
    HOT = "FF2F86"
    LIGHT = "FFF2F7"
    BORDER_HEX = "FFD0E2"
    DARK = "111111"
    GRAY_HEX = "6B7280"
    SOFT_GRAY = "F7F8FA"

    styles = doc.styles
    styles["Normal"].font.name = "맑은 고딕"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    styles["Normal"].font.size = Pt(10)
    styles["Normal"].font.color.rgb = RGBColor(17, 17, 17)

    def set_cell_shading(cell, fill: str):
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = tc_pr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            tc_pr.append(shd)
        shd.set(qn("w:fill"), fill)

    def set_cell_border(cell, color: str = "E5E7EB", sz: str = "6"):
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()
        borders = tc_pr.first_child_found_in("w:tcBorders")
        if borders is None:
            borders = OxmlElement("w:tcBorders")
            tc_pr.append(borders)
        for edge in ("top", "left", "bottom", "right"):
            tag = "w:{}".format(edge)
            element = borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                borders.append(element)
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), sz)
            element.set(qn("w:space"), "0")
            element.set(qn("w:color"), color)

    def format_run(run, size=None, bold=False, color=None):
        run.font.name = "맑은 고딕"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        if size:
            run.font.size = Pt(size)
        run.bold = bold
        if color:
            run.font.color.rgb = RGBColor.from_string(color)

    def add_hyperlink(paragraph, text: str, url: str):
        if not url:
            run = paragraph.add_run(text)
            format_run(run, 9, False, GRAY_HEX)
            return
        part = paragraph.part
        r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        new_run = OxmlElement("w:r")
        r_pr = OxmlElement("w:rPr")
        color = OxmlElement("w:color")
        color.set(qn("w:val"), HOT)
        r_pr.append(color)
        underline = OxmlElement("w:u")
        underline.set(qn("w:val"), "single")
        r_pr.append(underline)
        new_run.append(r_pr)
        text_el = OxmlElement("w:t")
        text_el.text = text
        new_run.append(text_el)
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)

    def add_label_value(paragraph, label: str, value: str, label_color=GRAY_HEX):
        r1 = paragraph.add_run(label)
        format_run(r1, 9, True, label_color)
        r2 = paragraph.add_run(str(value))
        format_run(r2, 10, True, DARK)

    def add_card_table(rows: int, cols: int, widths=None):
        table = doc.add_table(rows=rows, cols=cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                set_cell_border(cell, "E5E7EB", "6")
        if widths:
            for row in table.rows:
                for idx, width in enumerate(widths):
                    if idx < len(row.cells):
                        row.cells[idx].width = Cm(width)
        return table

    def safe_join(values):
        return " · ".join([str(v) for v in values if str(v).strip()])

    best_route = routes[0] if routes else None
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("💗 성북구 데이트 플래너")
    format_run(r, 22, True, HOT)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("조건 입력 결과를 바탕으로 만든 실사용용 데이트 일정표")
    format_run(r, 10, False, GRAY_HEX)

    doc.add_paragraph()

    # 1페이지 핵심 요약 카드
    summary_table = add_card_table(2, 4, widths=[4, 4, 4, 4])
    labels = ["지역", "분위기", "이동수단", "예산"]
    values = [
        meta.get("region", "-"),
        safe_join(meta.get("moods", [])),
        meta.get("transport", "-"),
        "{:,.0f}원".format(meta.get("budget", 0)),
    ]
    for i, label in enumerate(labels):
        set_cell_shading(summary_table.cell(0, i), LIGHT)
        p = summary_table.cell(0, i).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label)
        format_run(run, 9, True, HOT)
        p2 = summary_table.cell(1, i).paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p2.add_run(str(values[i]))
        format_run(run, 11, True, DARK)

    if best_route:
        doc.add_paragraph()
        course_line = " → ".join(p["name"] for p in best_route["places"])
        overview = add_card_table(2, 3, widths=[5, 5, 5])
        overview_labels = ["전체거리", "총 이동시간", "1인당 예상비용"]
        overview_values = ["{:.2f}km".format(best_route["distance_km"]), "약 {}분".format(best_route["minutes"]), "{:,}원".format(best_route["cost"])]
        for i, label in enumerate(overview_labels):
            set_cell_shading(overview.cell(0, i), SOFT_GRAY)
            p = overview.cell(0, i).paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(label)
            format_run(run, 9, True, GRAY_HEX)
            p2 = overview.cell(1, i).paragraphs[0]
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p2.add_run(overview_values[i])
            format_run(run, 15, True, HOT)

        p = doc.add_paragraph()
        r = p.add_run("오늘의 추천 동선")
        format_run(r, 13, True, DARK)
        p = doc.add_paragraph()
        r = p.add_run(course_line)
        format_run(r, 11, True, DARK)
        p = doc.add_paragraph()
        r = p.add_run("선택한 분위기와 예산, 이동 조건을 반영해 당일 바로 따라갈 수 있도록 구성한 코스입니다.")
        format_run(r, 9, False, GRAY_HEX)

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("이용 전 안내")
    format_run(r, 12, True, HOT)
    notes = [
        "평점·영업상태는 카카오 로컬 API에서 직접 제공하지 않아 카카오맵 링크에서 최종 확인하세요.",
        "예상비용은 업종별 평균 가격 모델을 기준으로 계산한 값입니다.",
        "날씨나 현장 상황에 따라 산책/야외 장소는 대체 장소를 확인하는 것을 추천합니다.",
    ]
    for note in notes:
        p = doc.add_paragraph(style=None)
        p.paragraph_format.left_indent = Cm(0.25)
        r = p.add_run("□ ")
        format_run(r, 9, True, HOT)
        r = p.add_run(note)
        format_run(r, 9, False, DARK)

    # 코스별 상세 페이지
    for idx, route in enumerate(routes, start=1):
        doc.add_page_break()
        h = doc.add_paragraph()
        r = h.add_run("코스 {}".format(idx))
        format_run(r, 18, True, HOT)
        r = h.add_run("  |  {}".format(" → ".join(p["name"] for p in route["places"])))
        format_run(r, 12, True, DARK)

        stat = add_card_table(2, 3, widths=[5, 5, 5])
        stat_labels = ["전체거리", "총 이동시간", "1인당 예상비용"]
        stat_values = ["{:.2f}km".format(route["distance_km"]), "약 {}분".format(route["minutes"]), "{:,}원".format(route["cost"])]
        for i, label in enumerate(stat_labels):
            set_cell_shading(stat.cell(0, i), LIGHT)
            pp = stat.cell(0, i).paragraphs[0]
            pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rr = pp.add_run(label)
            format_run(rr, 9, True, HOT)
            pp2 = stat.cell(1, i).paragraphs[0]
            pp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rr = pp2.add_run(stat_values[i])
            format_run(rr, 14, True, DARK)

        doc.add_paragraph()
        p = doc.add_paragraph()
        r = p.add_run("타임라인")
        format_run(r, 13, True, DARK)

        for order, place in enumerate(route["places"], start=1):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            r = p.add_run("{:02d}  ".format(order))
            format_run(r, 12, True, HOT)
            r = p.add_run(place["name"])
            format_run(r, 12, True, DARK)
            r = p.add_run("  ({})".format(place["type"]))
            format_run(r, 9, False, GRAY_HEX)

            detail = doc.add_paragraph()
            detail.paragraph_format.left_indent = Cm(0.75)
            add_label_value(detail, "예상비용: ", "{:,}원  ".format(place["cost"]))
            add_label_value(detail, "분위기: ", "{}  ".format(place.get("mood_tags", "-")))
            add_label_value(detail, "주소: ", place.get("address", "-"))
            link_p = doc.add_paragraph()
            link_p.paragraph_format.left_indent = Cm(0.75)
            r = link_p.add_run("카카오맵: ")
            format_run(r, 9, True, GRAY_HEX)
            add_hyperlink(link_p, "장소 열기", place.get("url", ""))

            if order <= len(route["segments"]):
                seg = route["segments"][order - 1]
                move = doc.add_paragraph()
                move.paragraph_format.left_indent = Cm(0.75)
                r = move.add_run("↓ {} 약 {}분 / {:.2f}km".format(seg["mode"], seg["minutes"], seg["distance_km"]))
                format_run(r, 9, True, HOT)
                r = move.add_run("   ")
                add_hyperlink(move, "길찾기", seg.get("url", ""))

        doc.add_paragraph()
        p = doc.add_paragraph()
        r = p.add_run("장소별 상세 카드")
        format_run(r, 13, True, DARK)

        for order, place in enumerate(route["places"], start=1):
            card = add_card_table(5, 2, widths=[3.0, 12.0])
            for row in card.rows:
                set_cell_border(row.cells[0], BORDER_HEX, "8")
                set_cell_border(row.cells[1], BORDER_HEX, "8")
            title_cell = card.cell(0, 0).merge(card.cell(0, 1))
            set_cell_shading(title_cell, LIGHT)
            title_p = title_cell.paragraphs[0]
            rr = title_p.add_run("📍 {}. {}".format(order, place["name"]))
            format_run(rr, 11, True, HOT)
            info = [
                ("분류", "{} / {}".format(place["type"], place["price_category"])),
                ("예상비용", "{:,}원".format(place["cost"])),
                ("분위기", place.get("mood_tags", "-")),
                ("주소", place.get("address", "-")),
            ]
            for r_idx, (label, value) in enumerate(info, start=1):
                set_cell_shading(card.cell(r_idx, 0), SOFT_GRAY)
                pp = card.cell(r_idx, 0).paragraphs[0]
                rr = pp.add_run(label)
                format_run(rr, 9, True, GRAY_HEX)
                pp = card.cell(r_idx, 1).paragraphs[0]
                rr = pp.add_run(str(value))
                format_run(rr, 9, False, DARK)

            reason = doc.add_paragraph()
            reason.paragraph_format.left_indent = Cm(0.25)
            rr = reason.add_run("추천 이유: ")
            format_run(rr, 9, True, HOT)
            rr = reason.add_run("선택한 분위기 키워드와 맞고, 코스 내 다른 장소와의 이동거리가 짧아 동선 효율이 좋습니다.")
            format_run(rr, 9, False, DARK)
            doc.add_paragraph()

        if route["segments"]:
            p = doc.add_paragraph()
            r = p.add_run("구간별 이동 정보")
            format_run(r, 13, True, DARK)
            table = add_card_table(1, 5, widths=[5.2, 2.3, 2.2, 2.2, 3.0])
            headers = ["구간", "이동수단", "거리", "예상시간", "길찾기"]
            for i, header in enumerate(headers):
                set_cell_shading(table.cell(0, i), LIGHT)
                pp = table.cell(0, i).paragraphs[0]
                rr = pp.add_run(header)
                format_run(rr, 9, True, HOT)
            for seg in route["segments"]:
                row = table.add_row().cells
                vals = ["{} → {}".format(seg["from"], seg["to"]), seg["mode"], "{:.2f}km".format(seg["distance_km"]), "약 {}분".format(seg["minutes"]), "길찾기"]
                for i, val in enumerate(vals):
                    pp = row[i].paragraphs[0]
                    if i == 4:
                        add_hyperlink(pp, val, seg.get("url", ""))
                    else:
                        rr = pp.add_run(str(val))
                        format_run(rr, 8.5, False, DARK)
                    set_cell_border(row[i], "E5E7EB", "6")

        doc.add_paragraph()
        p = doc.add_paragraph()
        r = p.add_run("예산 요약")
        format_run(r, 13, True, DARK)
        budget_table = add_card_table(1, 3, widths=[7, 4, 4])
        headers = ["항목", "분류", "예상비용"]
        for i, h in enumerate(headers):
            set_cell_shading(budget_table.cell(0, i), LIGHT)
            rr = budget_table.cell(0, i).paragraphs[0].add_run(h)
            format_run(rr, 9, True, HOT)
        for place in route["places"]:
            row = budget_table.add_row().cells
            vals = [place["name"], place["price_category"], "{:,}원".format(place["cost"])]
            for i, v in enumerate(vals):
                rr = row[i].paragraphs[0].add_run(str(v))
                format_run(rr, 9, False, DARK)
                set_cell_border(row[i], "E5E7EB", "6")
        total = budget_table.add_row().cells
        total[0].merge(total[1])
        set_cell_shading(total[0], SOFT_GRAY)
        rr = total[0].paragraphs[0].add_run("총 예상비용")
        format_run(rr, 10, True, DARK)
        set_cell_shading(total[2], LIGHT)
        rr = total[2].paragraphs[0].add_run("{:,}원".format(route["cost"]))
        format_run(rr, 11, True, HOT)

        doc.add_paragraph()
        p = doc.add_paragraph()
        r = p.add_run("데이트 전 체크리스트")
        format_run(r, 13, True, DARK)
        checklist = ["카카오맵에서 영업 여부 최종 확인", "예약 필요 여부 확인", "이동 경로 확인", "우천 시 대체 장소 확인", "예산 확인"]
        for item in checklist:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.25)
            r = p.add_run("□ ")
            format_run(r, 10, True, HOT)
            r = p.add_run(item)
            format_run(r, 10, False, DARK)

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


if "routes" not in st.session_state:
    st.session_state.routes = None
if "meta" not in st.session_state:
    st.session_state.meta = None
if "last_error" not in st.session_state:
    st.session_state.last_error = None

with st.sidebar:
    st.markdown("### 💗 조건 입력")
    with st.form("condition_form", clear_on_submit=False):
        region = st.selectbox("지역", list(REGION_KEYWORDS.keys()), index=2)
        moods = st.multiselect("분위기 키워드", ["감성", "조용한", "활동적", "힙한", "레트로", "가성비"], default=["감성", "조용한"])
        transport = st.selectbox("이동수단", ["도보", "대중교통", "대중교통+도보", "자동차"], index=0)
        budget = st.slider("1인 예산", min_value=0, max_value=100000, value=50000, step=10000)
        allow_over_budget = st.checkbox("예산 조금 초과 가능", value=False, help="체크를 풀면 입력하신 예산을 절대로 초과하지 않고 100% 한도 내에서만 매칭합니다.")
        place_count = st.selectbox("코스별 장소 수", [1, 2, 3, 4, 5], index=3)
        max_distance = st.slider("희망 전체거리 (km)", min_value=0.5, max_value=5.0, value=1.0, step=0.5, help="코스 전체 동선이 이 거리 이내인 코스만 추천합니다.")
        min_rating = st.slider("최소 평점", min_value=0.0, max_value=5.0, value=0.0, step=0.5)
        include_bar = st.checkbox("술집까지 포함", value=False)
        include_keyword = st.text_input("반드시 포함할 키워드", placeholder="예: 파스타, 전시, 인생네컷")
        submitted = st.form_submit_button("💗 새로운 코스 추천하기", use_container_width=True)
    clear_button = st.button("결과 지우기", use_container_width=True)

if clear_button:
    st.session_state.routes = None
    st.session_state.meta = None
    st.session_state.last_error = None

st.markdown(
    """
<div class="hero">
  <div class="hero-inner">
    <div class="hero-icon">💗</div>
    <div class="hero-copy">
      <div class="hero-kicker">Auto Mate · Smart Date Planner</div>
      <h1>성북구 <span class="hero-highlight">데이트 코스</span> 자동 추천</h1>
      <p>카카오 API 기반 실제 장소를 수집해 분위기·예산·동선 조건에 맞는 코스를 자동 생성합니다.</p>
    </div>
  </div>
</div>
<div class="flow">사용자 입력 → 카카오 API 장소 수집 → 후보 정제 → 키워드 매칭 → 구간별 이동시간 계산 → 지도 시각화 → DOCX 결과물 생성</div>
""",
    unsafe_allow_html=True,
)

if submitted:
    st.session_state.last_error = None
    api_key = get_api_key()
    if not api_key:
        st.session_state.last_error = "KAKAO_REST_API_KEY가 없습니다. 로컬은 .env, Streamlit Cloud는 Secrets 설정을 확인하세요."
    elif not moods:
        st.session_state.last_error = "분위기 키워드를 1개 이상 선택하세요."
    else:
        try:
            with st.spinner("카카오 API에서 장소를 수집하고 코스를 생성하는 중입니다..."):
                include_terms = parse_terms(include_keyword)
                df = collect_places(region, include_bar, api_key, include_terms)
                if df.empty:
                    raise RuntimeError("카카오 API에서 조건에 맞는 성북구 장소를 찾지 못했습니다.")
                df = add_scores(df, moods, include_terms)
                routes = build_routes(df, int(budget), transport, include_bar, int(place_count), float(max_distance), bool(allow_over_budget), include_terms)
                
                if not routes:
                    hint = ""
                    if include_bar:
                        hint = " 술집 포함 시 예상비용이 크게 올라가므로 예산을 높이거나 '예산 조금 초과 가능'을 체크해보세요."
                    if include_terms:
                        hint += f" 반드시 포함할 키워드({', '.join(include_terms)})가 너무 좁으면 결과가 줄어들 수 있습니다."
                    st.session_state.last_error = f"설정하신 예산({budget:,}원)·거리({max_distance:.1f}km) 조건에서 장소 {place_count}개 코스를 찾지 못했습니다." + hint
                    st.session_state.routes = None
                    st.session_state.meta = None
                else:
                    st.session_state.routes = routes
                    st.session_state.meta = {
                        "region": region,
                        "moods": list(moods),
                        "transport": transport,
                        "budget": int(budget),
                        "candidate_count": int(len(df)),
                        "include_bar": bool(include_bar),
                        "place_count": int(place_count),
                        "max_distance": float(max_distance),
                        "min_rating": float(min_rating),
                        "include_keyword": include_keyword,
                    }
        except Exception as exc:
            st.session_state.routes = None
            st.session_state.meta = None
            st.session_state.last_error = str(exc)

if st.session_state.last_error:
    st.error(st.session_state.last_error)

routes = st.session_state.routes
meta = st.session_state.meta

if routes and meta:
    c1, c2, c3 = st.columns(3)
    c1.markdown(f"<div class='metric-card'><div class='metric-title'>후보 장소</div><div class='metric-value'>{meta['candidate_count']}개</div></div>", unsafe_allow_html=True)
    c2.markdown(f"<div class='metric-card'><div class='metric-title'>선택 지역</div><div class='metric-value'>{meta['region']}</div></div>", unsafe_allow_html=True)
    c3.markdown(f"<div class='metric-card'><div class='metric-title'>코스별 장소 수</div><div class='metric-value'>{meta['place_count']}개</div></div>", unsafe_allow_html=True)

    sort_by = st.selectbox("정렬 기준", ["추천순", "총 이동시간 오름차순", "총 이동시간 내림차순", "1인당 예상비용 오름차순", "1인당 예상비용 내림차순", "전체거리 오름차순", "전체거리 내림차순"], index=0)
    shown_routes = sort_routes(routes, sort_by)

    for idx, route in enumerate(shown_routes, start=1):
        names = " → ".join(p["name"] for p in route["places"])
        budget_label = "예산 내" if route["budget_ok"] else "조건 조정 최적코스"
        st.markdown(
            f"""
<div class="route-card">
  <h2>코스 {idx} · ✅ {budget_label}</h2>
  <div class="stat-row">
    <div class="stat"><div class="label">전체거리</div><div class="num">{route['distance_km']:.2f}km</div></div>
    <div class="stat"><div class="label">총 이동시간</div><div class="num">약 {route['minutes']}분</div></div>
    <div class="stat"><div class="label">1인당 예상 비용</div><div class="num">{route['cost']:,}원</div></div>
  </div>
  <p style="margin-top:14px; font-weight:850; color:{BLACK};">{', '.join(meta['moods'])} 분위기에 맞춘 <b>{names}</b> 동선입니다.</p>
</div>
""",
    unsafe_allow_html=True,
)
        rows = []
        for order, p in enumerate(route["places"], start=1):
            rows.append({
                "순서": order,
                "장소명": p["name"],
                "분류": p["type"],
                "가격분류": p["price_category"],
                "1인당 예상비용": f"{p['cost']:,}원",
                "분위기": p["mood_tags"],
                "평점": p["rating"],
                "영업상태": p["open_status"],
                "주소": p["address"],
                "출처": p["source"],
            })
        st.dataframe(pd.DataFrame(rows), use_container_width=True, hide_index=True)

        if route["segments"]:
            st.markdown("**구간별 이동 정보**")
            seg_rows = []
            for s in route["segments"]:
                seg_rows.append({"구간": f"{s['from']} → {s['to']}", "이동수단": s["mode"], "거리": f"{s['distance_km']:.2f}km", "예상시간": f"약 {s['minutes']}분", "안내": s["guide"]})
            st.dataframe(pd.DataFrame(seg_rows), use_container_width=True, hide_index=True)

        st.markdown("**지도 미리보기**")
        st_folium(make_map(route), height=420, use_container_width=True, returned_objects=[])

        st.markdown("**장소 카카오맵 열기**")
        cols = st.columns(min(len(route["places"]), 5))
        for i, p in enumerate(route["places"]):
            with cols[i % len(cols)]:
                st.link_button(f"{i+1}. {p['name']} 열기", p["url"], use_container_width=True)

        if route["segments"]:
            st.markdown("**구간별 길찾기**")
            cols = st.columns(min(len(route["segments"]), 4))
            for i, s in enumerate(route["segments"]):
                with cols[i % len(cols)]:
                    st.link_button(f"{s['from']} → {s['to']} 길찾기", s["url"], use_container_width=True)

    st.download_button(
        "DOCX 데이트 플래너 다운로드",
        data=create_docx_bytes(shown_routes, meta),
        file_name="seongbuk_date_course_planner.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
else:
    if not st.session_state.last_error:
        st.info("왼쪽 조건을 선택한 뒤 '새로운 코스 추천하기' 버튼을 누르면 결과가 생성됩니다.")
